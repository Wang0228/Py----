import sys
sys.path.append(r'c:\users\user\appdata\local\programs\python\python311\lib\site-packages')
import os
from docx import Document

def read_word_file(filepath):
    doc=Document(filepath)#找到文件
    filename=os.path.basename(filepath)
    print("處理 \"",filename,"\" 文檔")
    full_text=[]
    for paragraph in doc.paragraphs:#doc.garagraph 將word拆成段落再加入到full_text
        full_text.append(paragraph.text)
    return '\n'.join(full_text)     #組合成一個字串回傳(整個word的字)

def Openai_ask(text,filepath,start=0,step=9900,overlap=3300):
    
    if(len(text)<overlap):
        overlap=0
    
    n=0
    while start<len(text)-overlap:#還有文本則繼續
        end=start+step
        doc=Document()
        doc.add_paragraph(text[start:end])  # 將文本添加到Word文檔中
        n+=1
        base_name, file_extension = os.path.splitext(os.path.basename(filepath))
        filename = f"{base_name}_{n}{file_extension}"
        file_path = os.path.join(r"C:\Users\user\Desktop\Py問題生成\AfterProcessing", filename)
        doc.save(file_path)  # 儲存Word文檔
        start=end-overlap#讓下一個迴圈從要overlap的地方開始
        

def main():
    folder_path = r"C:\Users\user\Desktop\Py問題生成\word"
    allfile = os.listdir(folder_path)
    for i in allfile:
        p = os.path.join(folder_path, i)
        a=read_word_file(p)
        Openai_ask(a,p)
        print(f"{i} 處理完成")
    

if __name__ == '__main__':
    main()